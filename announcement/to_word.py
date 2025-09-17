# Copyright 2025 Dagim Genene
# 
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# 
#     https://www.apache.org/licenses/LICENSE-2.0
# 
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.



import os
import django
import docx

# Setup Django environment if running as standalone script
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "chencha_school.settings")
django.setup()

from announcement.models import AutoSectionGrade

def write_to_word_section():
    document = docx.Document()
    students = AutoSectionGrade.objects.exclude(section__isnull=True).exclude(section='').values('section', 'student_name', 'father_name', 'grandfather_name','age', 'sex','grade').order_by('section','student_name')
    from collections import defaultdict


    grouped = defaultdict(list)
    for s in students:
        grouped[s["section"]].append(s)
        
    for section, students in grouped.items():
        heading = document.add_heading(f"{grouped[section][0]['grade']} - Section {section} students list",)
        run = heading.runs[0]
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)  # Black color
        table = document.add_table(rows=1, cols=6,style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Father Name'
        hdr_cells[3].text = 'Grandfather Name'
        hdr_cells[4].text = 'Sex'
        hdr_cells[5].text = 'Age'

        for i, s in enumerate(students,start=1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = s['student_name']
            row[2].text = s['father_name']
            row[3].text = s['grandfather_name']
            row[4].text = s['sex']
            row[5].text = str(s['age'])
        document.add_paragraph()  # Add a space between sections
    return document

def mark_list_in_word():
    document = docx.Document()

    students = AutoSectionGrade.objects.exclude(section__isnull=True).exclude(
        section=''
    ).values(
        'section', 'student_name', 'father_name', 'grandfather_name', 'age', 'sex', 'grade'
    ).order_by('section', 'student_name')

    from collections import defaultdict
    grouped = defaultdict(list)
    for s in students:
        grouped[s["section"]].append(s)

    for section, students in grouped.items():
        # Long heading line
        heading = document.add_heading(
            f"CHENCHA SECONDARY & PERPARATORY SCHOOL STUDENTS MARK LIST "
            f"YEAR 2018 E.C SEMESTER __________ GRADE & SECTION {grouped[section][0]['grade'][5:]} {section} "
            f"SUBJECT __________ SUBJECT TEACHER'S NAME __________ HOME ROOM TEACHER'S NAME __________",
            level=1
        )
        run = heading.runs[0]
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)  # Black color

        # Table with 14 columns
        table = document.add_table(rows=2, cols=14, style="Table Grid")

        # First header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[1].text = 'Name of Student'
        hdr_cells[2].text = 'Sex'
        hdr_cells[3].text = 'Age'
        hdr_cells[4].text = 'Test'
        hdr_cells[5].text = 'Lab/Other Activity'
        hdr_cells[6].text = 'Attendance'
        hdr_cells[7].text = 'Activity'
        hdr_cells[8].text = 'Total'
        hdr_cells[9].text = 'Mid'
        hdr_cells[10].text = 'Final'
        hdr_cells[11].text = 'Total'
        hdr_cells[12].text = 'Remarks'

        # Second header row with percentages
        perc_cells = table.rows[1].cells
        perc_cells[4].text = '10%'
        perc_cells[5].text = '10%'
        perc_cells[6].text = '5%'
        perc_cells[7].text = '5%'
        perc_cells[8].text = '30%'
        perc_cells[9].text = '20%'
        perc_cells[10].text = '50%'
        perc_cells[11].text = '100%'

        # Fill in student rows
        for i, s in enumerate(students, start=1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = f"{s['student_name']} {s['father_name']} {s['grandfather_name']}"
            row[2].text = s['sex']
            row[3].text = str(s['age'])
            # mark columns stay empty for teachers to fill

        document.add_paragraph()  # Space after each section

    return document
